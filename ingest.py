import os
import glob
from pathlib import Path
from typing import List, Any, Optional
from langchain_community.document_loaders import DirectoryLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document
from docx import Document as DocxDocument

from dotenv import load_dotenv

MODEL = "gpt-4.1-nano"

DB_NAME = str(Path(__file__).parent.parent / "vector_db")
KNOWLEDGE_BASE = str(Path(__file__).parent.parent / "knowledge-base")

# embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

load_dotenv(override=True)

embeddings = OpenAIEmbeddings(model="text-embedding-3-large")


def _load_docx_text(path: Path) -> str:
    """Extract plain text from a .docx file."""
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)


def _load_plain_text(path: Path) -> str:
    """Best-effort plain text loader for .txt / .md etc."""
    return path.read_text(encoding="utf-8", errors="ignore")


def fetch_documents_from_upload(uploaded_files: List[Any]) -> List[Document]:
    documents: List[Document] = []

    if not uploaded_files:
        return documents

    for f in uploaded_files:
        # Case 1: Gradio 6 NamedString (type="filepath")
        if hasattr(f, "data") and hasattr(f, "name") and not hasattr(f, "read"):
            path = Path(f.data)
            name = f.name

        # Case 2: plain string or Path
        elif isinstance(f, (str, Path)):
            path = Path(f)
            name = path.name

        # Case 3: file-like object with .read()
        else:
            raw_bytes = f.read()
            name = getattr(f, "name", "uploaded_file")
            suffix = Path(name).suffix.lower()

            if suffix == ".docx":
                # Need a filesystem path for python-docx; write temp if needed
                tmp_path = Path("tmp_upload.docx")
                tmp_path.write_bytes(raw_bytes)
                text = _load_docx_text(tmp_path)
                tmp_path.unlink(missing_ok=True)
            else:
                try:
                    text = raw_bytes.decode("utf-8")
                except UnicodeDecodeError:
                    text = raw_bytes.decode("latin-1")

            doc_type = suffix.lstrip(".") or "uploaded"
            documents.append(
                Document(
                    page_content=text,
                    metadata={"source": name, "doc_type": doc_type},
                )
            )
            continue  # go to next file

        # For the path-based cases (Gradio NamedString / str / Path)
        suffix = path.suffix.lower()

        if suffix == ".docx":
            text = _load_docx_text(path)
        else:
            # treat as plain text (txt/md) or at least try
            text = _load_plain_text(path)

        doc_type = suffix.lstrip(".") or "uploaded"

        documents.append(
            Document(
                page_content=text,
                metadata={"source": name, "doc_type": doc_type},
            )
        )

    return documents


def create_chunks(documents):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=200)
    chunks = text_splitter.split_documents(documents)
    return chunks


def create_embeddings(chunks):
    if os.path.exists(DB_NAME):
        Chroma(persist_directory=DB_NAME, embedding_function=embeddings).delete_collection()

    vectorstore = Chroma.from_documents(
        documents=chunks, embedding=embeddings, persist_directory=DB_NAME
    )

    collection = vectorstore._collection
    count = collection.count()

    sample_embedding = collection.get(limit=1, include=["embeddings"])["embeddings"][0]
    dimensions = len(sample_embedding)
    print(f"There are {count:,} vectors with {dimensions:,} dimensions in the vector store")
    return vectorstore


if __name__ == "__main__":
    documents = fetch_documents_from_upload()
    chunks = create_chunks(documents)
    print(f"[DEBUG] Created {len(chunks)} chunks")
    for i, ch in enumerate(chunks[:3]):
        print(f"\n[DEBUG] CHUNK {i} (len={len(ch.page_content)})")
        print(ch.page_content)
        print("-" * 80)
    create_embeddings(chunks)
    print("Ingestion complete")
